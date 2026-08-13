from pathlib import Path
import ipaddress
import tempfile
import unittest

from docx import Document

from pii_redaction import redact_docx
from pii_redaction.docx_processor import extract_docx_text
from pii_redaction.detectors import detect_entities
from pii_redaction.gold_fixture import GOLD_ENTITIES, NEGATIVE_STRINGS, run_gold_fixture


class RedactionTests(unittest.TestCase):
    def test_redacts_required_pii_and_keeps_plain_ids(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_path = tmp_path / "input.docx"
            output_path = tmp_path / "output.docx"
            document = Document()
            document.add_paragraph(
                "Rashi Patil emailed rashhi.patil@gmail.com and called +91 9876543210."
            )
            document.add_paragraph(
                "Company Acme Analytics Private Limited is at "
                "12 Green Park Road, New Delhi, Delhi 110016."
            )
            document.add_paragraph(
                "DOB: 14 May 1993. SSN: 123-45-6789. Card: 4111 1111 1111 1111. IP: 192.168.1.25."
            )
            document.add_paragraph("Keep these plain IDs: Ticket 77881 and Order 55773.")
            document.save(input_path)

            report = redact_docx(input_path, output_path)
            text = extract_docx_text(output_path)

            self.assertGreaterEqual(report["total_replacements"], 9)
            for original in [
                "Rashi Patil",
                "rashhi.patil@gmail.com",
                "+91 9876543210",
                "Acme Analytics Private Limited",
                "12 Green Park Road, New Delhi, Delhi 110016",
                "14 May 1993",
                "123-45-6789",
                "4111 1111 1111 1111",
                "192.168.1.25",
            ]:
                self.assertNotIn(original, text)
            self.assertIn("Ticket 77881", text)
            self.assertIn("Order 55773", text)

    def test_gold_fixture_covers_all_required_categories(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            score, _ = run_gold_fixture(Path(tmp))

            self.assertEqual(score["false_positive"], 0)
            self.assertEqual(score["false_negative"], 0)
            self.assertEqual(score["precision"], 1.0)
            self.assertEqual(score["recall"], 1.0)
            self.assertEqual(score["accuracy"], 1.0)
            self.assertEqual(
                set(score["per_type"]),
                {
                    "address",
                    "company",
                    "credit_card",
                    "dob",
                    "email",
                    "ip_address",
                    "person",
                    "phone",
                    "ssn",
                },
            )

    def test_structured_detector_negatives_are_not_redacted(self) -> None:
        text = (
            "Ticket 77881, Order 55773, invoice 2026-08, December 10, 2025, "
            "Section 32 of the Companies Act, 2013, invalid card 4111 1111 1111 1112, "
            "invalid IP 999.168.1.25, support line 12345, and addressed respectively."
        )
        self.assertEqual(detect_entities(text), [])

    def test_company_detector_ignores_generic_role_labels(self) -> None:
        text = (
            "Bankers to the Offer include Escrow Collection Bank, Refund Bank, "
            "Public Offer Account Bank and Sponsor Banks. Ratings mention Long Term Bank "
            "Facilities and Short Term Bank Facilities. RBI means Reserve Bank of India. "
            "The loan was with State Bank of India. Stock Exchanges include BSE Limited. "
            "Fragments such as India Limited, Pandit LLP, and Advisory Private Limited are not enough."
        )
        companies = [entity.text for entity in detect_entities(text) if entity.entity_type == "company"]
        for false_positive in [
            "Escrow Collection Bank",
            "Refund Bank",
            "Reserve Bank",
            "Short Term Bank",
            "State Bank",
            "India Limited",
            "Pandit LLP",
            "Advisory Private Limited",
            "BSE Limited",
        ]:
            self.assertNotIn(false_positive, companies)

    def test_company_detector_keeps_complete_legal_entity_spans(self) -> None:
        text = (
            "The book running lead managers are Nuvama Wealth Management Limited "
            "and ICICI Securities Limited. The registrar is MUFG Intime India Private Limited. "
            "The CARE Report was prepared by CARE Analytics and Advisory Private Limited. "
            "The auditors are Kirtane & Pandit LLP. "
            "SECI Solar Energy Corporation of India Limited appears in abbreviations. "
            "Suppliers include Vedanta Limited and Cindus Corporation."
        )
        companies = [entity.text for entity in detect_entities(text) if entity.entity_type == "company"]

        for expected in [
            "Nuvama Wealth Management Limited",
            "ICICI Securities Limited",
            "MUFG Intime India Private Limited",
            "CARE Analytics and Advisory Private Limited",
            "Kirtane & Pandit LLP",
            "Solar Energy Corporation of India Limited",
            "Vedanta Limited",
            "Cindus Corporation",
        ]:
            self.assertIn(expected, companies)
        self.assertNotIn("Advisory Private Limited", companies)
        self.assertNotIn("Pandit LLP", companies)
        self.assertNotIn("Solar Energy Corporation", companies)
        self.assertNotIn("SECI Solar Energy Corporation of India Limited", companies)

    def test_company_detector_does_not_merge_trust_list_with_company(self) -> None:
        text = (
            "Promoter Selling Shareholders include Kanchenjunga Family Trust "
            "and Waterloo Industrial Park VI Private Limited."
        )
        companies = [entity.text for entity in detect_entities(text) if entity.entity_type == "company"]

        self.assertIn("Waterloo Industrial Park VI Private Limited", companies)
        self.assertNotIn(
            "Kanchenjunga Family Trust and Waterloo Industrial Park VI Private Limited",
            companies,
        )

    def test_company_replacement_preserves_complete_span(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_path = tmp_path / "input.docx"
            output_path = tmp_path / "output.docx"
            document = Document()
            document.add_paragraph(
                "Prepared by CARE Analytics and Advisory Private Limited for review."
            )
            document.add_paragraph(
                "Certified by Kirtane & Pandit LLP, Chartered Accountants."
            )
            document.save(input_path)

            report = redact_docx(input_path, output_path)
            text = extract_docx_text(output_path)
            originals = {
                item["original"]
                for item in report["entities"]
                if item["type"] == "company"
            }

            self.assertIn("CARE Analytics and Advisory Private Limited", originals)
            self.assertIn("Kirtane & Pandit LLP", originals)
            self.assertNotIn("Advisory Private Limited", originals)
            self.assertNotIn("Pandit LLP", originals)
            self.assertNotIn("CARE Analytics and Advisory Private Limited", text)
            self.assertNotIn("Kirtane & Pandit LLP", text)
            self.assertNotIn("CARE Analytics and", text)

    def test_redacts_grouped_indian_phone_number(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_path = tmp_path / "input.docx"
            output_path = tmp_path / "output.docx"
            document = Document()
            document.add_paragraph("Phone: +91 98765 43210")
            document.add_paragraph("Repeat: +91 98765 43210")
            document.save(input_path)

            report = redact_docx(input_path, output_path)
            text = extract_docx_text(output_path)
            phones = [item for item in report["mapping"] if item["type"] == "phone"]

            self.assertNotIn("+91 98765 43210", text)
            self.assertEqual(phones[0]["original"], "+91 98765 43210")
            self.assertRegex(phones[0]["replacement"], r"^\+91 \d{5} \d{5}$")

    def test_ipv6_addresses_are_single_valid_replacements(self) -> None:
        samples = [
            "2001:db8:85a3::8a2e:370:7334",
            "2001:0db8:85a3:0000:0000:8a2e:0370:7334",
        ]
        for sample in samples:
            with self.subTest(sample=sample):
                entities = detect_entities(f"IPv6: {sample}.")
                self.assertEqual(len(entities), 1)
                self.assertEqual(entities[0].text, sample)
                self.assertEqual(entities[0].entity_type, "ip_address")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_path = tmp_path / "input.docx"
            output_path = tmp_path / "output.docx"
            document = Document()
            for sample in samples:
                document.add_paragraph(f"IPv6: {sample}")
            document.save(input_path)

            report = redact_docx(input_path, output_path)
            text = extract_docx_text(output_path)
            mapping = {
                item["original"]: item["replacement"]
                for item in report["mapping"]
                if item["type"] == "ip_address"
            }

            for sample in samples:
                self.assertNotIn(sample, text)
                replacement = mapping[sample]
                self.assertEqual(ipaddress.ip_address(replacement).version, 6)
                self.assertIn(replacement, text)

    def test_mixed_text_address_is_redacted_as_complete_span(self) -> None:
        address = "42 Lake View Road, Sector 18, Chandigarh, 160018, India"
        paragraph = (
            f"The company office is at {address}. "
            "The server used for testing is 192.168.10.24."
        )
        entities = detect_entities(paragraph)
        addresses = [entity for entity in entities if entity.entity_type == "address"]
        self.assertEqual(len(addresses), 1)
        self.assertEqual(addresses[0].text, address)

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_path = tmp_path / "input.docx"
            output_path = tmp_path / "output.docx"
            document = Document()
            document.add_paragraph(paragraph)
            document.save(input_path)

            redact_docx(input_path, output_path)
            text = extract_docx_text(output_path)

            self.assertNotIn(address, text)
            self.assertNotIn("Chandigarh, 160018, India", text)
            self.assertNotIn("41100118", text)

    def test_repeated_pii_uses_consistent_fake_replacement(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_path = tmp_path / "input.docx"
            output_path = tmp_path / "output.docx"
            document = Document()
            document.add_paragraph("Rashi Patil emailed rashhi.patil@gmail.com.")
            document.add_paragraph("Rashi Patil followed up from rashhi.patil@gmail.com.")
            document.save(input_path)

            report = redact_docx(input_path, output_path)
            text = extract_docx_text(output_path)
            mapping = {
                (item["type"], item["original"]): item["replacement"]
                for item in report["mapping"]
            }

            self.assertNotIn("Rashi Patil", text)
            self.assertNotIn("rashhi.patil@gmail.com", text)
            self.assertEqual(text.count(mapping[("person", "Rashi Patil")]), 2)
            self.assertEqual(text.count(mapping[("email", "rashhi.patil@gmail.com")]), 2)

    def test_gold_fixture_negatives_survive_redaction(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            _, run_report = run_gold_fixture(Path(tmp))
            redacted_text = extract_docx_text(Path(tmp) / "gold_fixture_redacted.docx")

            self.assertGreaterEqual(len(GOLD_ENTITIES), 19)
            self.assertEqual(run_report["residual_original_values"], [])
            for value in NEGATIVE_STRINGS:
                self.assertIn(value, redacted_text)


if __name__ == "__main__":
    unittest.main()
