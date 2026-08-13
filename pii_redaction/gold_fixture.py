"""Controlled gold fixture used for defensible evaluation metrics."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from . import redact_docx
from .docx_processor import extract_docx_text
from .evaluation import score_gold_fixture


GOLD_ENTITIES = [
    {"type": "person", "value": "Rashi Patil"},
    {"type": "person", "value": "Rohan Dey"},
    {"type": "person", "value": "Dr Meera Nair"},
    {"type": "email", "value": "rashhi.patil@gmail.com"},
    {"type": "email", "value": "rohan.dey@gmail.com"},
    {"type": "phone", "value": "+91 9876543210"},
    {"type": "phone", "value": "+91 98765 43210"},
    {"type": "phone", "value": "987-654-3210"},
    {"type": "company", "value": "Acme Analytics Private Limited"},
    {"type": "company", "value": "Zenith Data Services LLP"},
    {"type": "address", "value": "12 Green Park Road, New Delhi, Delhi 110016"},
    {"type": "address", "value": "44 Maple Street, Springfield, IL 62704"},
    {"type": "address", "value": "42 Lake View Road, Sector 18, Chandigarh, 160018, India"},
    {"type": "ssn", "value": "123-45-6789"},
    {"type": "ssn", "value": "987-65-4321"},
    {"type": "credit_card", "value": "4111 1111 1111 1111"},
    {"type": "credit_card", "value": "5555-5555-5555-4444"},
    {"type": "dob", "value": "14 May 1993"},
    {"type": "dob", "value": "03/08/1988"},
    {"type": "ip_address", "value": "192.168.1.25"},
    {"type": "ip_address", "value": "2001:0db8:85a3:0000:0000:8a2e:0370:7334"},
    {"type": "ip_address", "value": "2001:db8:85a3::8a2e:370:7334"},
]

NEGATIVE_STRINGS = [
    "Ticket 77881",
    "Order 55773",
    "invoice 2026-08",
    "December 10, 2025",
    "Section 32 of the Companies Act, 2013",
    "4111 1111 1111 1112",
    "999.168.1.25",
    "support line 12345",
    "addressed respectively",
    "Order ID: ORD-2026-48172",
    "Ticket ID: TKT-88421",
    "Meeting date: 13 August 2026",
    "Version: 2.4.1",
]


def build_gold_fixture(path: str | Path) -> None:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_paragraph(
        "Rashi Patil opened ticket 10452 from rashhi.patil@gmail.com. "
        "Rohan Dey later wrote from rohan.dey@gmail.com and called from +91 9876543210."
    )
    document.add_paragraph(
        "Alternate Indian mobile format: +91 98765 43210."
    )
    document.add_paragraph(
        "Dr Meera Nair joined from Zenith Data Services LLP and uses office phone 987-654-3210."
    )
    document.add_paragraph(
        "Employer: Acme Analytics Private Limited. Address: "
        "12 Green Park Road, New Delhi, Delhi 110016. Mailing address: "
        "44 Maple Street, Springfield, IL 62704."
    )
    document.add_paragraph(
        "SSN 123-45-6789 and backup SSN 987-65-4321. Cards: "
        "4111 1111 1111 1111 and 5555-5555-5555-4444."
    )
    document.add_paragraph(
        "DOB: 14 May 1993. Date of birth: 03/08/1988. "
        "IPs: 192.168.1.25 and 2001:0db8:85a3:0000:0000:8a2e:0370:7334."
    )
    document.add_paragraph(
        "Mixed office sentence: The company office is at "
        "42 Lake View Road, Sector 18, Chandigarh, 160018, India. "
        "Compressed IPv6: 2001:db8:85a3::8a2e:370:7334."
    )
    document.add_paragraph(
        "Negatives to keep: Ticket 77881, Order 55773, invoice 2026-08, "
        "December 10, 2025, Section 32 of the Companies Act, 2013, "
        "invalid card 4111 1111 1111 1112, invalid IP 999.168.1.25, "
        "support line 12345, addressed respectively, Order ID: ORD-2026-48172, "
        "Ticket ID: TKT-88421, Meeting date: 13 August 2026, and Version: 2.4.1."
    )
    document.save(path)


def run_gold_fixture(output_dir: str | Path) -> tuple[dict, dict]:
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    fixture_path = output_dir / "gold_fixture.docx"
    redacted_path = output_dir / "gold_fixture_redacted.docx"
    report_json = output_dir / "gold_fixture_run.json"

    build_gold_fixture(fixture_path)
    run_report = redact_docx(fixture_path, redacted_path, report_json)
    predicted = [{"type": item["type"], "value": item["original"]} for item in run_report["entities"]]
    score = score_gold_fixture(predicted, GOLD_ENTITIES)

    redacted_text = extract_docx_text(redacted_path)
    for item in GOLD_ENTITIES:
        if item["value"] in redacted_text:
            raise AssertionError(f"Gold PII value leaked after redaction: {item['value']}")
    for value in NEGATIVE_STRINGS:
        if value not in redacted_text:
            raise AssertionError(f"Negative fixture value was unexpectedly redacted: {value}")

    return score, run_report
