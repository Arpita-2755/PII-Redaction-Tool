"""Evaluation report helpers for redaction runs."""

from __future__ import annotations

from collections import Counter
from pathlib import Path
import json
from typing import Iterable


def score_gold_fixture(predicted: Iterable[dict], gold: Iterable[dict]) -> dict:
    predicted_set = {(item["type"], item["value"]) for item in predicted}
    gold_set = {(item["type"], item["value"]) for item in gold}
    true_positive = len(predicted_set & gold_set)
    false_positive = len(predicted_set - gold_set)
    false_negative = len(gold_set - predicted_set)
    precision = _safe_div(true_positive, true_positive + false_positive)
    recall = _safe_div(true_positive, true_positive + false_negative)
    accuracy = _safe_div(true_positive, true_positive + false_positive + false_negative)
    per_type = {}
    for entity_type in sorted({item[0] for item in predicted_set | gold_set}):
        predicted_type = {item for item in predicted_set if item[0] == entity_type}
        gold_type = {item for item in gold_set if item[0] == entity_type}
        type_tp = len(predicted_type & gold_type)
        type_fp = len(predicted_type - gold_type)
        type_fn = len(gold_type - predicted_type)
        per_type[entity_type] = {
            "true_positive": type_tp,
            "false_positive": type_fp,
            "false_negative": type_fn,
            "precision": round(_safe_div(type_tp, type_tp + type_fp), 4),
            "recall": round(_safe_div(type_tp, type_tp + type_fn), 4),
            "accuracy": round(_safe_div(type_tp, type_tp + type_fp + type_fn), 4),
        }
    return {
        "true_positive": true_positive,
        "false_positive": false_positive,
        "false_negative": false_negative,
        "precision": round(precision, 4),
        "recall": round(recall, 4),
        "accuracy": round(accuracy, 4),
        "per_type": per_type,
    }


def write_markdown_report(
    run_report: dict,
    output_path: str | Path,
    gold_score: dict | None = None,
) -> None:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    counts = Counter(run_report.get("counts_by_type", {}))
    total = run_report.get("total_replacements", 0)
    residuals = run_report.get("residual_original_values", [])
    replacement_rate = run_report.get("detected_candidate_replacement_rate", 1.0)

    lines = [
        "# Evaluation Strategy and Metrics",
        "",
        "## Strategy",
        "",
        "The supplied Red Herring Prospectus does not include a human-annotated ground-truth label file, so I used a two-part evaluation:",
        "",
        "1. A controlled gold fixture with known PII labels and known non-PII negatives. This is the only source used for true precision, recall, and accuracy.",
        "2. A prospectus run audit that records every detected candidate, verifies that detected originals no longer remain in the output, checks repeated-value consistency, and reports per-type replacement counts.",
        "",
        "Structured PII is detected with strict regular expressions and validators. Credit cards must pass Luhn validation. Dates of birth are redacted only near DOB context words to avoid treating normal prospectus dates as birth dates. Names, companies, and addresses use conservative heuristics, so the main tradeoff is better precision with some possible missed uncommon names.",
        "",
    ]

    if gold_score:
        lines.extend(
            [
                "## Gold Fixture Metrics",
                "",
                "These are entity-level set metrics measured against the controlled fixture ground truth. They are not claimed as real-world benchmark accuracy.",
                "",
                f"- True positives: {gold_score['true_positive']}",
                f"- False positives: {gold_score['false_positive']}",
                f"- False negatives: {gold_score['false_negative']}",
                f"- Precision: {gold_score['precision']:.2%}",
                f"- Recall: {gold_score['recall']:.2%}",
                f"- Accuracy: {gold_score['accuracy']:.2%}",
                "",
                "| PII Type | TP | FP | FN | Precision | Recall | Accuracy |",
                "| --- | ---: | ---: | ---: | ---: | ---: | ---: |",
            ]
        )
        for entity_type, score in sorted(gold_score.get("per_type", {}).items()):
            lines.append(
                f"| {entity_type} | {score['true_positive']} | {score['false_positive']} | "
                f"{score['false_negative']} | {score['precision']:.2%} | "
                f"{score['recall']:.2%} | {score['accuracy']:.2%} |"
            )
        lines.append("")
    else:
        lines.extend(
            [
                "## Gold Fixture Metrics",
                "",
                "No gold fixture score was supplied for this report, so true precision, recall, and accuracy are not reported here.",
                "",
            ]
        )

    lines.extend(
        [
        "## Prospectus Audit",
        "",
        f"- Total replacements: {total}",
        f"- XML parts scanned: {run_report.get('xml_parts_scanned', 0)}",
        f"- Paragraphs scanned: {run_report.get('paragraphs_scanned', 0)}",
        f"- Residual original PII values found after replacement: {len(residuals)}",
        f"- Detected-candidate replacement rate: {replacement_rate:.2%}",
        "- Prospectus precision/recall/accuracy: not claimed because no human-labeled prospectus ground truth was provided.",
        "",
        "## Counts by PII Type",
        "",
        "| PII Type | Replacements |",
        "| --- | ---: |",
        ]
    )
    for entity_type, count in sorted(counts.items()):
        lines.append(f"| {entity_type} | {count} |")

    lines.extend(
        [
            "",
            "## False Positive / False Negative Notes",
            "",
            "- Ticket IDs, order numbers, page numbers, and normal financial figures are not redacted unless they match a protected PII type.",
            "- Company names with clear legal suffixes are redacted; generic product or section names are left intact.",
            "- Dates are treated as DOB only when DOB wording is nearby. This improves precision for a prospectus, which contains many business dates.",
            "- Some rare person names without honorifics or first names outside the built-in list may require adding a custom rule or dictionary.",
            "",
            "## Residual Check",
            "",
        ]
    )
    if residuals:
        lines.append("Residual originals were found and should be reviewed:")
        for residual in residuals[:25]:
            lines.append(f"- {residual['type']}: `{residual['value']}`")
    else:
        lines.append("No original detected PII values were found in the generated output DOCX.")

    lines.append("")
    output_path.write_text("\n".join(lines), encoding="utf-8")


def write_docx_report(markdown_path: str | Path, docx_path: str | Path) -> None:
    """Create a simple DOCX report from the generated Markdown."""

    from docx import Document
    from docx.shared import Inches, Pt

    markdown_path = Path(markdown_path)
    docx_path = Path(docx_path)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    for raw_line in markdown_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=1)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|") and "---" not in line:
            document.add_paragraph(line.replace("|", "  ").strip())
        elif line.startswith("1. ") or line.startswith("2. "):
            document.add_paragraph(line[3:], style="List Number")
        else:
            document.add_paragraph(line)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


def load_report(path: str | Path) -> dict:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def _safe_div(numerator: int | float, denominator: int | float) -> float:
    return float(numerator) / denominator if denominator else 0.0
